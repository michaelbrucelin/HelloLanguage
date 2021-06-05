# 这里使用python-docx来操作word文档，注意是python-docx，而不是docx，docx是另一个模块，但是python-docx导入的时候使用的是import docx
# python3 -m pip install python-docx

import win32com.client
import docx

# 1. 读取word
doc = docx.Document('data/demo.docx')
len(doc.paragraphs)             # 7
doc.paragraphs[0].text          # 'Document Title'
doc.paragraphs[1].text  # 'A plain paragraph with some bold and some italic'
len(doc.paragraphs[1].runs)     # 5
doc.paragraphs[1].runs[0].text  # 'A plain paragraph with'
doc.paragraphs[1].runs[1].text  # ' some '
doc.paragraphs[1].runs[2].text  # 'bold'
doc.paragraphs[1].runs[3].text  # ' and some '
doc.paragraphs[1].runs[4].text  # 'italic'

# 2. 从word中获取全部文本


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        # fullText.append('    '+para.text)  # 每一个段落有缩进
    return '\n'.join(fullText)
    # return '\n\n'.join(fullText)  # 段落之间增加空行


print(getText('data/demo.docx'))

# 3. 更改word的样式
doc = docx.Document('data/demo.docx')
doc.paragraphs[0].text   # 'Document Title'
doc.paragraphs[0].style  # _ParagraphStyle('Title') id: 58339184
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[0].style  # _ParagraphStyle('Normal') id: 58424096
doc.paragraphs[1].text   # 'A plain paragraph with some bold and some italic'
(doc.paragraphs[1].runs[0].text,
 doc.paragraphs[1].runs[1].text,
 doc.paragraphs[1].runs[2].text,
 doc.paragraphs[1].runs[3].text,
 doc.paragraphs[1].runs[4].text)
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[2].underline = True
doc.paragraphs[1].runs[4].underline = True
doc.save('data/restyled.docx')

# 4. 创建文档
doc = docx.Document()
doc.add_paragraph('Hello, world!')
doc.save('data/helloworld.docx')

doc = docx.Document()
doc.add_paragraph('Hello, world!', 'Title')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('data/multipleParagraphs.docx')

# 4. 添加标题
doc = docx.Document()
doc.add_heading('Head 0', 0)
doc.add_heading('Head 0', 1)
doc.add_heading('Head 0', 2)
doc.add_heading('Head 0', 3)
doc.add_heading('Head 0', 4)
doc.save('data/heading.docx')

# 5. 添加换行符与换页符
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('data/twoPage.docx')

# 6. 添加图像
doc = docx.Document()
doc.add_picture('data/zophie.png',
                width=docx.shared.Inches(1),
                height=docx.shared.Cm(4))
doc.save('data/pic.docx')

# 7. 使用word创建pdf
# PyPDF2不支持创建pdf，但是如果在Windows上安装了Microsoft Word，那么结合使用pywin32模块可以将word转为pdf
# pip install pywin32
# import win32com.client

wordFilename = 'data/twoPage.docx'
pdfFilename = 'data/docxto.pdf'

wdFormatPDF = 17
wordObj = win32com.client.Dispatch('Word.Application')
docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.close()
wordObj.close()
